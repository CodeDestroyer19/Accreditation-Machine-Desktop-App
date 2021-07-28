import docx
import os

AllfileList = {}

TemplateDoc = docx.Document('tester.docx')

path = '/Users/mikek/Documents/doc23/docs/Docs/AFRI-Umalusi/Case 3'
dirs = os.listdir(path)

if os.path.exists(path+'/QuickRun/save'):
    print('It exisits')
else:
    os.mkdir(path+'/QuickRun/save')

# To fill Template document

doc_cont = {}


def Populate(key, value):
    paraLib = {}
    doc = docx.Document(value)
    i = 0
    for para in doc.paragraphs:
        i += 1
        if (len(para.text) > 0):
            paraLib[i] = para.text
    return paraLib

# to get Docs whose content will be copied to empty doc


def map():
    savePath = os.listdir(os.path.join(path + '/QuickRun/save'))
    try:
        for key, value in list(doc_cont.items()):
            if key in savePath:
                del doc_cont[key]
                return doc_cont
            else:
                for keys, values in list(value.items()):
                    TemplateDoc.add_paragraph(values)
                    TemplateDoc.save(os.path.join(
                        path + '/QuickRun/save', key))
                    value.pop(keys)
            del doc_cont[key]

    except Exception as e:
        print(e)


def FilerComp():
    i = 0
    for key, value in AllfileList.items():
        doc_cont[key] = Populate(key, value)
        map()


for dir in dirs:
    is_dir = os.path.isdir(os.path.join(path, dir))

    if is_dir == True:
        folder = os.path.join(path, dir).__str__()
        itemsInfolder = os.listdir(folder)

        for file in itemsInfolder:
            if file.endswith('.docx'):
                AllfileList[file] = os.path.join(folder, file)


FilerComp()


# TemplateDoc.add_paragraph(para.text)

# for table in document.tables:
#     tables = TemplateDoc.add_table(
#         len(table.table.rows), len(table.table.columns))
#     for i in range(len(table.rows)):
#         for row in table.rows:
#             for cell in row.cells:
#                 print(i)
#                 tables.add_row().cells[i].text = cell.text
#                 break
#             break
#     break
