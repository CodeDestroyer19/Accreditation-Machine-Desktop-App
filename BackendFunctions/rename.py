import os
import shutil
import time
from PySide2 import QtCore
from pdf2docx import Converter
import win32com.client
import threading
from docx import Document


class Rename(object):

    def RenameFolder(newName, renameStr, path, dir):
        if not os.path.isdir(os.path.join(path, dir)):
            pass
        if newName in dir:
            pass
        os.rename(os.path.join(path, dir), os.path.join(
            path, dir.replace(renameStr, newName)))

    def ConvertMultiFileToDoc(self, pathOffile, doc_type):
        word = win32com.client.Dispatch("Word.application")

        def callback():

            for file in pathOffile:

                index = file.find('.')
                docx_file = '{0}{1}'.format(file[:index], '.docx')
                if doc_type == '.pdf':
                    try:
                        cv = Converter(file)
                        cv.convert(docx_file, start=0, end=None)
                        cv.close()
                        self.ui.outPutConvertMultiFile.addItem(file)

                    except Exception as e:
                        print(e)
                elif doc_type == '.doc':
                    try:
                        wordDoc = word.Documents.Open(file[2:])
                        wordDoc.SaveAs(docx_file, FileFormat=16)
                        wordDoc.Close()
                        self.ui.outPutConvertMultiFile.addItem(file)

                    except Exception as e:
                        print('Failed to Convert: {0}'.format(file))
                        print(e)
                else:
                    document = Document()
                    print(document)
                    myfile = open(file, encoding="utf-8").read()
                    document.add_paragraph(myfile)
                    document.save(docx_file)
                    self.ui.outPutConvertMultiFile.addItem(file)

            time.sleep(10)
            self.ui.dirInfoMulti.setText('No file selected')
            self.ui.outPutConvertMultiFile.clear()
            self.ui.outPutConvertMultiFile.addItem(
                'Select other files to convert!')

        t = threading.Thread(target=callback)
        t.start()

    def ConvertSingleFileToDoc(self, pathOffile, doc_type):
        word = win32com.client.Dispatch("Word.application")
        index = pathOffile.find('.')
        docx_file = '{0}{1}'.format(pathOffile[:index], '.docx')

        def callback():
            if doc_type == '.pdf':
                try:
                    cv = Converter(pathOffile)
                    cv.convert(docx_file, start=0, end=None)
                    cv.close()
                    self.ui.outputLabelConvertSingle.setText('Complete!')
                    time.sleep(5)
                    self.ui.dirInfo.setText('No file selected')
                    self.ui.outputLabelConvertSingle.setText(
                        'Select another file to convert!')

                except Exception as e:
                    print(e)
            elif doc_type == '.doc':
                paths = pathOffile[2:]
                try:
                    wordDoc = word.Documents.Open(paths)
                    wordDoc.SaveAs(docx_file, FileFormat=16)
                    wordDoc.Close()
                    self.ui.outputLabelConvertSingle.setText('Complete!')
                    time.sleep(5)
                    self.ui.dirInfo.setText('No file selected')
                    self.ui.outputLabelConvertSingle.setText(
                        'Select another file to convert!')

                except Exception as e:
                    print('Failed to Convert: {0}'.format(pathOffile))
                    print(e)
            else:
                document = Document()
                myfile = open(pathOffile, encoding="utf-8").read()
                document.add_paragraph(myfile)
                document.save(docx_file)
                self.ui.outputLabelConvertSingle.setText('Complete!')
                time.sleep(5)
                self.ui.dirInfo.setText('No file selected')
                self.ui.outputLabelConvertSingle.setText(
                    'Select other file to convert!')

        t = threading.Thread(target=callback)
        t.start()

    def RenameFile(self, newName, renameStr, path, file):
        folderName = self.ui.OutFolderName.text()
        if self.outpath == None:
            pass
        else:
            Folder = os.path.join(self.outpath, folderName)
            if renameStr in file:

                file_path = os.path.join(path, file)

                file_dist_same_path = os.path.join(
                    path, file.replace(renameStr, newName))

                if len(folderName) > 0:
                    if os.path.exists(Folder):
                        pass
                    else:
                        os.mkdir(Folder)

                    file_dist_custom_outpath = os.path.join(Folder, file)

                    try:
                        shutil.copy(file_path, file_dist_custom_outpath)
                        os.rename(os.path.join(Folder, file), os.path.join(
                            Folder, file.replace(renameStr, newName)))
                    except Exception as e:
                        self.ui.OutPutDisplay.addItem(e)
                else:
                    os.rename(file_path, file_dist_same_path)

                self.ui.OutPutDisplay.addItem('Done Renaming: ' + file)

    def SumitPath(self, Dirs):
        path = self.path
        actualPath = path.replace('You have selected: ', '')

        if Dirs == True:
            Rename.runParseForFile(self, actualPath)
        elif Dirs == False:
            Rename.runParseForFilesInFolder(self, actualPath)
        else:
            pass

    def toggleMkdirOption(self, maxheight, enable):
        if enable:
            height = self.ui.OutFolderName.height()
            maxExtend = maxheight
            standard = 0

            if height == 0:
                heightExtended = maxExtend
            else:
                heightExtended = standard

            self.animation = QtCore.QPropertyAnimation(
                self.ui.OutFolderName, b"minimumHeight")
            self.animation.setDuration(400)
            self.animation.setStartValue(height)
            self.animation.setEndValue(heightExtended)
            self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuad)
            self.animation.start()

    def runParseForFile(self, path):
        self.ui.FolderItemDisplay.clear()
        renameStr = self.ui.RepalceValueInputRenamePage.text()
        newName = self.ui.NewValueInputRenamePage.text()

        files = os.listdir(path[2:])

        for index, file in enumerate(files):
            if os.path.isfile(os.path.join(path, file)) == True:
                if file.endswith('.docx') or file.endswith('.doc'):
                    pathOffile = os.path.join(path, file)
                    docx_file = '{0}{1}'.format(pathOffile, 'x')

                    Rename.RenameFile(
                        self, newName, renameStr, path, file)
                # RenameFolder('NewNAme', 'OldName', path, dir)
                # ConvertDoc(file, pathOffile, docx_file)

    def ItemsBoxFiles(self, path):
        filesholder = {}
        files = os.listdir(path)
        for index, file in enumerate(files):
            is_file = os.path.isfile(os.path.join(path, file))
            if is_file == True:
                if file.endswith('.docx') or file.endswith('.doc'):
                    filesholder[index] = file
                    self.ui.FolderItemDisplay.isWrapping = True

                    if len(filesholder) > 0:
                        self.ui.LenItemsRenamePage.setText(
                            len(filesholder.items()).__str__() + ' items found')
                    if filesholder.items() == None:
                        self.ui.LenItemsRenamePage.setText(
                            '0 items found')

                    self.ui.FolderItemDisplay.addItem(file)

    def ItemsBoxFolders(self, path):
        filesholder = {}
        dirs = os.listdir(path)
        for dir in dirs:
            is_dir = os.path.isdir(os.path.join(path, dir))
            if is_dir == True:
                folder = os.path.join(path, dir).__str__()
                itemsInfolder = os.listdir(folder)
                for index, file in enumerate(itemsInfolder):
                    is_file = os.path.isfile(os.path.join(folder, file))
                    if is_file == True:
                        if file.endswith('.docx') or file.endswith('.doc'):
                            filesholder[index] = file
                            self.ui.FolderItemDisplay.isWrapping = True
                            if len(filesholder) > 0:
                                self.ui.LenItemsRenamePage.setText(
                                    len(filesholder.items()).__str__() + ' items found')
                            if filesholder.items() == None:
                                self.ui.LenItemsRenamePage.setText(
                                    '0 items found')

                            self.ui.FolderItemDisplay.addItem(file)

    def runParseForFilesInFolder(self, path):
        renameStr = self.ui.RepalceValueInputRenamePage.text()
        newName = self.ui.NewValueInputRenamePage.text()

        dirs = os.listdir(path[2:])
        for dir in dirs:
            is_dir = os.path.isdir(os.path.join(path, dir))

            if is_dir == True:
                folder = os.path.join(path, dir).__str__()
                itemsInfolder = os.listdir(folder)

                for index, file in enumerate(itemsInfolder):
                    is_file = os.path.isfile(os.path.join(folder, file))
                    if is_file == True:
                        if file.endswith('.docx') or file.endswith('.doc'):
                            pathOffile = os.path.join(folder, file)
                            docx_file = '{0}{1}'.format(pathOffile, 'x')

                            Rename.RenameFile(
                                self, newName, renameStr, folder, file)
                        # RenameFolder('NewNAme', 'OldName', path, dir)
                        # ConvertDoc(file, pathOffile, docx_file)
