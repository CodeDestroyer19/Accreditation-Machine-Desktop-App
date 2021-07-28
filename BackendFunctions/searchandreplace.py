import re
import docx
import os

path = '/Users/mikek/Documents/doc23/docs/Docs/AFRI-Umalusi/Case 3'
dirs = os.listdir(path)

renameStr = input('What do you want to replace? Type here then hit enter... ')
newName = input('What are you replacing it with? Type here then hit enter... ')


def docx_replace_regex(doc_obj, regex, replace, targetText):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs

    for paragraph in doc_obj.paragraphs:
        if targetText in paragraph.text:
            print(paragraph.text)
            paragraph.text = paragraph.text.replace(targetText, replace)
        else:
            print('No matches')


for dir in dirs:
    is_dir = os.path.isdir(os.path.join(path, dir))

    if is_dir == True:
        folder = os.path.join(path, dir).__str__()
        itemsInfolder = os.listdir(folder)

        for file in itemsInfolder:
            is_file = os.path.isfile(os.path.join(folder, file))

            if is_file == True:
                if 'docx' in file:
                    regex1 = re.compile(r"your regex")
                    doc = docx.Document(os.path.join(folder, file))
                    docx_replace_regex(doc, regex1, newName, renameStr)
                    doc.save(os.path.join(folder, file))
                else:
                    print('No results')