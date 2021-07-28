from BackendFunctions.audit import Audit
import docx
import os
from docx.shared import Cm, Pt


class PopulateDocs(object):

    def docx_replace_img(doc_obj):
        section = doc_obj.sections[0]
        section.different_first_page_header_footer = True

        for sec in doc_obj.sections:
            sec.right_margin = Cm(1.91)
            sec.left_margin = Cm(1.91)

        for shape in section.header.paragraphs:
            if len(shape.text) > 0:
                print(len(shape.text))
                shape.runs[0].clear()

        for texts in section.footer.paragraphs:
            texts.clear()
            if len(texts.text) > 0:
                print(len(texts.text))
                texts.text = ''

        paraText = section.header.paragraphs[0]
        paraText.clear()
        paragraph_format = paraText.paragraph_format
        paragraph_format.space_before = Pt(24)
        paraText.text = 'Testing Something'

        section.header_distance = Cm(0)
        header = section.first_page_header
        header.paragraphs[0].clear()
        paraIMG = header.paragraphs[0]
        paraIMG.paragraph_format.left_indent = Cm(-1.9)
        img = paraIMG.add_run()
        # img.add_picture(None, width=Cm(21.60))

    def iter(self, files, path):
        for file in files:
            filePath = os.path.join(path, file)
            if os.path.isfile(filePath) == True:
                if file.endswith('.docx'):
                    doc = docx.Document(filePath)

        for file in files:
            if file.endswith('.docx'):
                pathtoFile = os.path.join(path, file)
                DisplayDoc = docx.Document(pathtoFile)
                print(DisplayDoc._element)
                self.ui.WordDisplay.setControl(DisplayDoc._element)
                self.ui.WordDisplay.show()
            break

    def ParseToplevelDir(self, path):

        finalpath = path[2:]

        files = os.listdir(finalpath)
        PopulateDocs.iter(self, files, finalpath)

    def ParseNestedDir(self, path):

        finalpath = path[2:]
        dirs = os.listdir(finalpath)

        for dir in dirs:
            folder = os.path.join(finalpath, dir)
            if os.path.isdir(folder) == True:
                files = os.listdir(folder)
                PopulateDocs.iter(self, files, folder)

    def triggerDoc(self):
        self.ui.WordDisplay.setControl(
            '/Users/mikek/Documents/quinn kiwalabye kiw001 event 3 swatchbook.docx')
        self.ui.WordDisplay.show()

    #                 pathtoFile = os.path.join(folder, file)
    #                 doc = docx.Document(pathtoFile)
    #                 docx_replace_img(doc)
    #                 print('\nDone appling Template to: ', file)
    #                 doc.save(os.path.join(testPathFinal, file))
    #                 print('Performing Audit on: ', file, '\n')
    #                 savePath = os.path.join(testPathFinal, file)
    #                 docAudit = docx.Document(savePath)
    #                 Audit(docAudit, savePath)

# MeDA@#2020
